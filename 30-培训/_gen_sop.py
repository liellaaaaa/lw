# -*- coding: utf-8 -*-
"""生成 WorkBuddy 通用配置与使用 SOP（第一层·最低使用基线）Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_FONT = "微软雅黑"
TITLE_FONT = "微软雅黑"

doc = Document()

# ---------- 基础字体（中英文统一）----------
def set_run_font(run, name=BASE_FONT, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color

# 设置 Normal 样式默认字体
normal = doc.styles['Normal']
normal.font.name = BASE_FONT
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), BASE_FONT)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.25

# 页边距
for s in doc.sections:
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)

# ---------- 辅助函数 ----------
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        set_run_font(r, name=TITLE_FONT, size=15 if level == 1 else (13 if level == 2 else 11.5),
                     bold=True, color=RGBColor(0x1F, 0x3B, 0x73))
    return h

def para(text="", bold=False, italic=False, size=10.5, color=None, align=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, bold=bold, italic=italic, size=size, color=color)
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, size=10.5)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True, size=10.5)
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    return p

def screenshot_placeholder(label):
    """插入一个带浅灰底色的占位框，提示此处需放截图"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    # 浅灰底
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EEF1F6')
    tcPr.append(shd)
    cell.width = Cm(15)
    # 文本
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("【截图占位 · " + label + "】")
    set_run_font(r, bold=True, size=10, color=RGBColor(0x55, 0x5F, 0x70), italic=True)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("（此处放置对应界面截图；截图需包含操作入口与关键按钮，下附一句话说明）")
    set_run_font(r2, size=9, color=RGBColor(0x80, 0x88, 0x96), italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def info_table(rows, col_widths=None, header=None):
    ncol = len(rows[0])
    t = doc.add_table(rows=0, cols=ncol)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header:
        hr = t.add_row().cells
        for i, h in enumerate(header):
            hr[i].text = ""
            p = hr[i].paragraphs[0]
            r = p.add_run(h)
            set_run_font(r, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
            # 蓝底
            tcPr = hr[i]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1F3B73')
            tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================================================
# 封面
# ============================================================
sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(60)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("WorkBuddy 通用配置与使用 SOP")
set_run_font(r, name=TITLE_FONT, size=24, bold=True, color=RGBColor(0x1F, 0x3B, 0x73))
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("第一层 · 最低使用基线（消除非业务性使用障碍）")
set_run_font(r, size=13, bold=True, color=RGBColor(0x44, 0x55, 0x66))
sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("统一账号权限、基础配置、任务发起、资料调用、结果检查与异常反馈")
set_run_font(r, size=10.5, color=RGBColor(0x66, 0x6F, 0x7A))
doc.add_paragraph()

# 文档信息表
info_table(
    [
        ["文档名称", "WorkBuddy 通用配置与使用 SOP（第一层）"],
        ["版本 / 状态", "V1.0 · 草案（待截图补全）"],
        ["适用对象", "公司全员（首次接触 WorkBuddy 的同事）"],
        ["编写", "AI 技术专员（肖聪）"],
        ["日期", "2026-08-11"],
        ["配套资料", "《WorkBuddy 培训索引》、各岗位实操手册（后续）"],
    ],
    col_widths=[4, 11],
)

# ============================================================
# 0. 文档定位与边界
# ============================================================
heading("0. 文档定位与边界说明", 1)
para("本文档是 WorkBuddy 培训体系的第一层底座，解决的核心问题不是「让所有人的能力一致」，而是先消除一批非业务性的使用障碍——配置不同、权限不清、资料调用方式不同、结果不会检查、遇到问题不会反馈。", space_after=6)
para("通过统一一批最低使用动作（账号、权限、基础配置、任务发起、资料调用、结果检查、异常反馈），把因环境和操作差异造成的「用不起来 / 用错」降到最低，让所有人先站到同一条起跑线上。", space_after=6)

heading("0.1 关键边界：基础动作可统一，方法不必一致", 2)
para("不要把「统一最低使用基线」理解成「要求所有人用一模一样的方法」。下列内容可以且应当统一：", bold=False, space_after=2)
bullet("统一的入口与配置（账号、权限、项目、基础设置）；")
bullet("统一的动作名称与最小步骤（任务怎么发起、资料怎么给、结果怎么查）；")
bullet("统一的异常反馈格式（问题 / 截图 / 输入 / 预期）。")
para("但以下内容应保留岗位差异，由各场景自行定义，不在本 SOP 强制统一：", bold=False, space_after=2)
bullet("具体提示词（prompt）的写法与话术；")
bullet("任务拆解的流程与顺序；")
bullet("结果「好 / 不好」的判断标准（因岗位目标而异）。")

heading("0.2 培训落地路径", 2)
para("通用配置与使用 SOP → 统一基础任务练习 → 岗位真实场景实操 → 结果验收与问题台账。", bold=True, space_after=2)
bullet("第一步（本文档）：统一最低使用基线，能正常配置、能发起任务、能调用资料、会检查与反馈。")
bullet("第二步：用一组通用的基础任务练习，验证「会操作工具」。")
bullet("第三步：回到各岗位真实业务场景实操，推动从「会操作」走向「能独立完成一项可验收的工作任务」。")
bullet("第四步：建立结果验收标准与问题台账，持续收集失败案例并反哺本 SOP。")

# ============================================================
# 1. 账号、权限与基础配置
# ============================================================
heading("1. 账号、权限与基础配置", 1)
para("目标：让每位同事进入 WorkBuddy 后，环境一致、权限清晰、工作文件有固定落点，避免因「我的界面和你不一样」而卡住。", space_after=6)

heading("1.1 账号登录与激活", 2)
numbered("使用公司分配的账号登录 WorkBuddy（桌面客户端或网页端）。", )
numbered("首次登录后确认头像 / 昵称正确，并完成必要的初始化引导。")
numbered("如登录异常（无响应、反复退出），先检查网络与客户端版本，必要时联系 AI 技术专员。")
screenshot_placeholder("登录入口与登录后主界面")

heading("1.2 权限说明（先讲清，避免越权或不敢用）", 2)
info_table(
    [
        ["能力", "默认权限", "说明"],
        ["发起对话 / 使用技能", "全员可用", "基础使用不受限"],
        ["读取本机文件 / 工作区", "仅当前项目文件夹", "AI 只应访问你打开的项目目录"],
        ["连接器（企微 / 邮箱等）", "需单独授权 / Trust", "连接外部系统须本人授权，状态以侧栏为准"],
        ["外发动作（发消息 / 邮件）", "谨慎，先确认", "涉及对外发送需人工确认，不在本层默认放开"],
        ["删除 / 覆盖文件", "高风险，需确认", "批量或不可逆操作前必须二次确认"],
    ],
    col_widths=[4.5, 4, 6.5],
    header=["能力", "默认权限", "说明"],
)

heading("1.3 新建文件夹并设置为「项目」（工作文件落点）", 2)
para("把工作文件放到固定目录，是后续所有资料调用、结果保存、复用能跑通的前提。请每位同事为自己的常用工作建一个项目文件夹。", space_after=4)
numbered("在文件管理器或 WorkBuddy 工作区侧栏中，新建一个文件夹（建议命名含部门 / 用途，如「采购部-日常」）。")
numbered("在 WorkBuddy 中将该文件夹打开 / 设为当前「项目（工作区）」。")
numbered("确认后，AI 生成的工作文件默认落在此文件夹内；项目内会生成 .workbuddy 子目录用于存储配置，请勿删除。")
numbered("后续调用资料、保存结果都基于这个文件夹，避免文件散落桌面 / 下载目录。")
screenshot_placeholder("新建文件夹并设为当前项目/工作区")
screenshot_placeholder("项目文件夹内 .workbuddy 目录（勿删）提示")

heading("1.4 专家 / 技能 / 连接器：入口与开启", 2)
para("这三类是 WorkBuddy 的能力扩展入口，统一在左侧栏或对应管理页进入。本层只要求「知道入口、会开启、看懂状态」，具体用法在岗位实操阶段展开。", space_after=4)

para("（1）专家（Expert）", bold=True, space_after=2)
bullet("入口：左侧栏「专家」选项，按类别浏览并进入对话。", )
bullet("用途：召唤特定角色（如数据分析师）做专项任务；也可直接对话使用。")
screenshot_placeholder("专家中心入口与分类浏览界面")

para("（2）技能（Skill）", bold=True, space_after=2)
bullet("入口：技能管理页，可从推荐市场安装，或本地编写后安装。", )
bullet("存储位置：用户级在 ~/.workbuddy/skills/，项目级在 {工作区}/.workbuddy/skills/。", )
bullet("调用：对话框输入「/」可查看本机已装技能并触发。")
screenshot_placeholder("技能管理页（安装/启用/已装列表）")

para("（3）连接器（Connector）", bold=True, space_after=2)
bullet("入口：连接器管理页，列出所有可连外部系统（企业微信、邮箱、网盘、腾讯文档等）。", )
bullet("开启：添加配置后，须在管理页点击「信任 / Trust」才会真正启用；侧栏状态显示 connected / disconnected。", )
bullet("注意：连接外部系统需本人授权（如邮箱用授权码而非登录密码），AI 不能代替你完成授权。")
screenshot_placeholder("连接器管理页（状态与 Trust 按钮）")

heading("1.5 设置：智能体设置 与 个性化设置", 2)
para("统一两处基础设置，避免「同一条任务，不同人结果差异来自配置」。", space_after=4)
bullet("智能体设置（Agent / 模型设置）：确认当前使用的模型与「思考模式」开关状态；团队给出推荐默认值，个人可按机器性能微调。", bold_prefix="智能体设置 —— ")
bullet("个性化设置：填写必要的个人偏好（如语言、默认工作目录、回复风格）；让 AI 的默认行为符合你的习惯。", bold_prefix="个性化设置 —— ")
screenshot_placeholder("设置页-智能体设置")
screenshot_placeholder("设置页-个性化设置")

# ============================================================
# 2. 任务如何发起
# ============================================================
heading("2. 任务如何发起（背景与要求说明）", 1)
para("目标：让 AI 一次听懂你要做什么，减少来回拉扯。统一「先说背景、再说要求、最后给资料/约束」的表达顺序。", space_after=6)

heading("2.1 发起入口", 2)
bullet("在对话框直接输入任务；复杂任务建议先在心里拆成「背景 + 要求 + 资料 + 约束」四块。")
bullet("需要特定能力时，先召唤专家或触发技能，再给任务。")
screenshot_placeholder("对话发起入口 / 召唤专家或技能的位置")

heading("2.2 背景说明（让 AI 知道「为什么做」）", 2)
para("背景决定 AI 的判断口径。至少说清：这件事属于哪个业务、给谁用、要解决什么、有无时间或口径限制。", space_after=4)
para("示例：", bold=True, space_after=2)
para("「我是采购部同事，需要把本周原料价格波动整理成一份给主管的简报，用于周一例会汇报，数据以生意社为准。」", italic=True, color=RGBColor(0x33, 0x33, 0x33))

heading("2.3 要求说明（让 AI 知道「做到什么程度」）", 2)
para("用结构化方式写要求，比一句话更稳。推荐模板：", space_after=4)
info_table(
    [
        ["要素", "写法示例", "作用"],
        ["目标", "产出一份周度价格波动简报", "明确交付物"],
        ["范围", "仅氯化铵、硫酸铵等 5 个品种", "划定边界"],
        ["格式", "Markdown 表格 + 3 条结论", "统一形态"],
        ["约束", "数据来源限生意社，不含推测", "控制质量"],
        ["验收", "主管能直接用于例会汇报", "定义好坏"],
    ],
    col_widths=[2.5, 7, 5.5],
    header=["要素", "写法示例", "作用"],
)
para("提示：要求越具体，返工越少；但「具体」不等于「长」，把关键约束写清楚即可。岗位阶段的判断标准（什么算好）由各场景自定。", space_after=6)

# ============================================================
# 3. 资料如何调用
# ============================================================
heading("3. 资料如何调用（文件 / 知识库 / Skill / 连接器）", 1)
para("目标：让 AI 用对你的资料，而不是凭空编。统一「资料从哪来、怎么给、给到什么程度」的基线动作。", space_after=6)

heading("3.1 本地文件与文件夹", 2)
bullet("把要处理的文件放进第 1.3 节设置的项目文件夹，再在对话中说明文件名 / 路径。")
bullet("多个文件可说明「项目文件夹内某子目录下的全部文件」。")
bullet("敏感文件勿放公共项目；不确定时先问 AI 技术专员。")

heading("3.2 知识库（如 Obsidian 本地库）", 2)
bullet("知识库即项目文件夹内的 Markdown 文件，AI 可直接读取并引用。")
bullet("调用时说明「读 XX 文件夹 / XX 文件」，或「基于知识库最近 N 天的日报」。")
bullet("保持知识库文件命名规范，AI 才能稳定检索。")

heading("3.3 技能 Skill", 2)
bullet("重复型任务先装好对应技能，再用「/技能名」触发，避免每次重新描述流程。")
bullet("技能装好后，把「一次跑通的过程」沉淀为可复用指令。")

heading("3.4 连接器 Connector", 2)
bullet("已连接并 Trust 的连接器，可在任务中直接读写其内容（如读企业微信消息、读邮箱附件）。")
bullet("调用前确认连接器状态为 connected；断开时 AI 无法访问该外部系统。")
bullet("外部系统的「写 / 发」动作需本人确认，不在本层默认自动执行。")
screenshot_placeholder("在对话中引用项目文件/知识库/触发技能的示意")

# ============================================================
# 4. 输出结果如何检查、保存、复用
# ============================================================
heading("4. 输出结果如何检查、保存、复用", 1)
para("目标：不做「AI 说什么就是什么」的甩手掌柜。统一最低检查动作与保存习惯。", space_after=6)

heading("4.1 结果检查（最低动作）", 2)
bullet("看结构：交付物是否齐全（表格 / 结论 / 文件都齐了没）。")
bullet("看事实：关键数字、文件名、引用是否来自你给的资料，而非凭空生成。")
bullet("看口径：是否符合你第 2.3 节提的约束（来源、范围、格式）。")
bullet("存疑就问：对不确定的结论，要求 AI 标注来源或重新核对，不要默认采信。")
para("注：什么叫「结果可用」的具体标准由岗位场景定，但「会做以上四步检查」是本层统一要求。", italic=True, color=RGBColor(0x55, 0x5F, 0x70))

heading("4.2 保存位置", 2)
bullet("结果默认落在第 1.3 节的项目文件夹；重要交付物按业务归类子目录存放。")
bullet("Word / PPT / 报表等成品，保存到对应项目资料目录，便于他人复用与追溯。")
bullet("不要把成品只留在聊天框里——聊天记录不等于交付物。")

heading("4.3 复用方式", 2)
bullet("把跑通的任务沉淀为技能或知识库笔记，下次直接调用。")
bullet("周期性任务配置为自动化（如每日价格摘要），减少重复发起。")
screenshot_placeholder("结果文件在项目文件夹中的保存位置示意")

# ============================================================
# 5. 敏感数据 / 业务承诺 / 人工复核边界
# ============================================================
heading("5. 敏感数据、业务承诺与人工复核边界", 1)
para("目标：划清「AI 能做什么、不能替你担什么责」，防止越界使用。", space_after=6)

heading("5.1 敏感数据边界", 2)
bullet("个人信息、薪资、合同金额、客户名单等敏感数据，避免直接粘贴到对话，确需处理时放入私有项目并确认无外发。")
bullet("涉及对外系统的写入（发邮件、发企微、改文档），默认需本人确认后再执行。")

heading("5.2 业务承诺边界", 2)
bullet("AI 产出不得直接作为对客 / 对外的正式承诺（报价、交期、合同条款等）。")
bullet("任何对客正式内容，须由责任人复核签字后再发出。")

heading("5.3 人工复核边界", 2)
bullet("事实性结论、数字、法律 / 财务相关判断，必须人工复核。")
bullet("AI 是辅助，最终责任在使用人及其主管，不在工具。")
info_table(
    [
        ["场景", "AI 可做的", "必须人工复核"],
        ["价格波动简报", "汇总、制表、写初稿", "结论是否用于决策"],
        ["对客报价单", "生成草稿、核对格式", "金额、条款、交期"],
        ["合同 / 文书", "起草、改写", "法律风险、正式发出"],
        ["外发消息", "拟稿", "发送动作本身"],
    ],
    col_widths=[3.5, 6, 5.5],
    header=["场景", "AI 可做的", "必须人工复核"],
)

# ============================================================
# 6. 异常如何反馈
# ============================================================
heading("6. 异常如何反馈（问题描述 / 截图 / 输入材料 / 预期结果）", 1)
para("目标：遇到失败不要只说「用不了」，按统一格式反馈，让问题可复现、可定位、可解决。", space_after=6)

heading("6.1 反馈四要素（统一模板）", 2)
info_table(
    [
        ["要素", "要写什么", "示例"],
        ["问题描述", "具体现象，而非「不行」", "触发技能后无响应，卡在加载"],
        ["截图", "报错界面 / 异常状态", "见下方占位（含入口与按钮）"],
        ["输入材料", "你给了什么（文件/指令）", "已上传存货统计表.xlsx 并输入…"],
        ["预期结果", "你认为该怎样", "应生成 5 个品种的价格简报"],
    ],
    col_widths=[2.5, 6, 6.5],
    header=["要素", "要写什么", "示例"],
)
screenshot_placeholder("反馈时附带的报错/异常界面截图")

heading("6.2 反馈渠道", 2)
bullet("日常问题：在工作群 @ AI 技术专员，附上述四要素。")
bullet("配置 / 权限类：直接找 AI 技术专员处理，勿自行改全局配置。")
bullet("重复出现的问题：记入「问题台账」，纳入下一轮 SOP 修订。")

# ============================================================
# 7. 培训有效性判断
# ============================================================
heading("7. 培训有效性判断（不只看「是否听懂 SOP」）", 1)
para("判断一个人是否真正用起来了，至少看三点，而非仅看课堂上是否听懂：", space_after=4)
numbered("能否独立完成基础任务——不依赖手把手，自己配好环境、发起并完成一个最小任务。", bold_prefix="① ")
numbered("能否判断结果是否可用——对 AI 产出做第 4.1 节的检查，能指出哪里不对。", bold_prefix="② ")
numbered("遇到失败时能否准确反馈问题——按第 6 节格式给出可复现的问题描述。", bold_prefix="③ ")
para("三点都满足，才算「会操作工具」；再进入岗位真实场景，看能否独立完成一项可验收的工作任务。", bold=True, space_after=6)

# ============================================================
# 8. 附录
# ============================================================
heading("附录 A：基础任务练习清单（自检表）", 1)
info_table(
    [
        ["#", "基础任务", "完成标准", "是否掌握"],
        ["1", "登录并确认主界面", "能看到侧栏与对话入口", "□"],
        ["2", "新建文件夹并设为项目", "AI 文件落在该文件夹", "□"],
        ["3", "开启一个连接器并 Trust", "侧栏状态为 connected", "□"],
        ["4", "发起一个带背景+要求的任务", "AI 一次返回可用初稿", "□"],
        ["5", "调用一份本地文件/知识库", "AI 引用了你给的资料", "□"],
        ["6", "检查并保存结果", "文件落在项目目录", "□"],
        ["7", "按四要素反馈一个异常", "问题描述可复现", "□"],
    ],
    col_widths=[1, 5.5, 6.5, 2],
    header=["#", "基础任务", "完成标准", "是否掌握"],
)

heading("附录 B：术语对照", 1)
info_table(
    [
        ["术语", "含义"],
        ["项目 / 工作区", "你打开让 AI 读写文件的本地文件夹，工作成果落点"],
        ["专家 Expert", "可召唤的专项角色，做特定类型任务"],
        ["技能 Skill", "把一次跑通的任务固化成可复用指令"],
        ["连接器 Connector", "连接外部系统（企微/邮箱/网盘等），需授权与 Trust"],
        ["自动化 Automation", "定时运行某个技能并产出报告"],
        ["Trust", "连接器启用前需点击的信任确认"],
    ],
    col_widths=[3.5, 11.5],
    header=["术语", "含义"],
)

para()
para("— 文档结束 · V1.0 草案 · 截图占位待补全 —", italic=True, size=9, color=RGBColor(0x88, 0x88, 0x88), align=WD_ALIGN_PARAGRAPH.CENTER)

out = r"E:\lw\30-培训\WorkBuddy通用配置与使用SOP（第一层·最低使用基线）.docx"
doc.save(out)
print("SAVED:", out)
